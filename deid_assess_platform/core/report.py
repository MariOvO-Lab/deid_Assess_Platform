from docx import Document
from docx.shared import Inches
import datetime
import os

class ReportGenerator:
    def __init__(self):
        self.template_path = "templates/report_template.docx"
    
    def generate(self, evaluation_results):
        """
        生成评估报告
        :param evaluation_results: 评估结果
        :return: 报告文件路径
        """
        # 创建文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('去标识化及匿名化评估报告', 0)
        
        # 添加报告生成时间
        doc.add_paragraph(f'报告生成时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # 添加评估结果部分
        doc.add_heading('评估结果', level=1)
        
        # 添加评估指标表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '评估指标'
        hdr_cells[1].text = '结果'
        
        # 处理特殊情况
        smallest_classes = evaluation_results.pop('前5个最小等价类', None)
        violating_classes = evaluation_results.pop('违规等价类', None)
        
        for key, value in evaluation_results.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # 添加前5个最小等价类信息
        if smallest_classes:
            doc.add_heading('前5个最小等价类', level=2)
            # 使用表格显示
            class_table = doc.add_table(rows=1, cols=3)
            class_hdr_cells = class_table.rows[0].cells
            class_hdr_cells[0].text = '排名'
            class_hdr_cells[1].text = '大小'
            class_hdr_cells[2].text = '值'
            
            for i, cls in enumerate(smallest_classes):
                row_cells = class_table.add_row().cells
                row_cells[0].text = str(i+1)
                row_cells[1].text = str(cls['size'])
                row_cells[2].text = str(cls['values'])
        
        # 添加违规等价类信息
        if violating_classes:
            doc.add_heading('违规等价类（不满足K值）', level=2)
            # 使用表格显示
            violating_table = doc.add_table(rows=1, cols=3)
            violating_hdr_cells = violating_table.rows[0].cells
            violating_hdr_cells[0].text = '排名'
            violating_hdr_cells[1].text = '大小'
            violating_hdr_cells[2].text = '值'
            
            for i, cls in enumerate(violating_classes):
                row_cells = violating_table.add_row().cells
                row_cells[0].text = str(i+1)
                row_cells[1].text = str(cls['size'])
                row_cells[2].text = str(cls['values'])
        
        # 添加结论部分
        doc.add_heading('结论', level=1)
        
        # 基于评估结果生成结论
        # 获取重识别风险（新的指标名称）
        reid_risk = float(evaluation_results.get('重识别风险上界', '0'))
        # 转换为0-100的范围以便于等级判断
        reid_risk *= 100
        
        # 获取可用性损失（已经是百分比字符串）
        availability_loss_str = evaluation_results.get('可用性损失', '0%')
        availability_loss = float(availability_loss_str.replace('%', ''))
        
        if reid_risk < 10:
            risk_level = '低'
        elif reid_risk < 30:
            risk_level = '中'
        else:
            risk_level = '高'
        
        if availability_loss < 30:
            availability_level = '高'
        elif availability_loss < 70:
            availability_level = '中'
        else:
            availability_level = '低'
        
        conclusion = f"本次去标识化处理后的重识别风险等级为{risk_level}，数据可用性等级为{availability_level}。"
        doc.add_paragraph(conclusion)
        
        # 保存报告
        report_dir = "reports"
        if not os.path.exists(report_dir):
            os.makedirs(report_dir)
        
        report_filename = f"{report_dir}/deid_evaluation_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(report_filename)
        
        return report_filename
